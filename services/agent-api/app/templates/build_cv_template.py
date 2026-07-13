"""Regenerate `cv_template.docx` — the fixed docxtpl template for the tailored CV.

The template is a checked-in binary; this script is its readable source of record.
Every placeholder here binds 1:1 to a field on `TailoredCV` (app/cv_tailor.py). If
you rename a pydantic field, rename its placeholder here and regenerate in the same
commit — otherwise the render regression test (tests/test_cv_render.py) fails on the
now-undefined variable.

Run from the agent-api directory:

    uv run python app/templates/build_cv_template.py

This is a dev/tooling script, not runtime code — app/cv_render.py only ever *reads*
the produced .docx; it never imports this module.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt


def build() -> Path:
    doc = Document()

    # Header block — profile identity + who this copy was tailored for + contact.
    # The "Tailored for" line is conditional: the tailored CV sets target_role, the
    # full CV (build_full_cv) leaves it empty so the line is dropped entirely.
    doc.add_heading("{{ profile_name }}", level=0)
    doc.add_paragraph("{{ profile_tagline }}")
    doc.add_paragraph("{%p if target_role %}")
    doc.add_paragraph("Tailored for: {{ target_role }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("{{ contact_email }} | {{ contact_linkedin }}")

    doc.add_heading("Summary", level=1)
    doc.add_paragraph("{{ generated_summary }}")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph("{{ skills|join(', ') }}")

    doc.add_heading("Experience", level=1)
    # `{%p ... %}` is docxtpl's paragraph tag: the paragraph holding it is removed
    # at render, so no blank lines are left behind by the loop scaffolding.
    doc.add_paragraph("{%p for role in roles %}")
    doc.add_heading("{{ role.title }} — {{ role.company }}", level=2)
    doc.add_paragraph("{{ role.date_range }}")
    doc.add_paragraph("{%p for hl in role.highlights %}")
    doc.add_paragraph("{{ hl.text }}", style="List Bullet")
    doc.add_paragraph("{%p endfor %}")
    doc.add_paragraph("Stack: {{ role.stack|join(', ') }}")
    doc.add_paragraph("{%p endfor %}")

    # Keep the default body font modest and predictable.
    style = doc.styles["Normal"]
    style.font.size = Pt(11)

    out = Path(__file__).with_name("cv_template.docx")
    doc.save(str(out))
    return out


if __name__ == "__main__":
    path = build()
    print(f"wrote {path}")
