"""Render regression tests for the CV-tailoring step 2 (app/cv_render.py).

These render a fixture TailoredCV against the *real* checked-in template and assert
a genuine, openable .docx comes out with every contract field present. The point is
to catch placeholder/field-name drift between the pydantic contract and the template
at build time (a failing test) instead of at runtime (a blank or broken CV).
"""
import os
from pathlib import Path

import jinja2
import pytest
from docx import Document
from docxtpl import DocxTemplate

from app.cv_render import TEMPLATE_PATH, render_cv
from app.cv_tailor import TailoredCV, TailoredHighlight, TailoredRole, build_full_cv


def _fixture_cv() -> TailoredCV:
    return TailoredCV(
        target_role="Staff Engineer",
        generated_summary="Architect & manager with 17+ years.",
        roles=[
            TailoredRole(
                id="smartlinx", title="Hands-On Manager & Architect", company="SmartLinx",
                date_range="Jan 2025 – Jun 2026", stack=["C#", "Azure"],
                highlights=[
                    TailoredHighlight(text="Re-platformed a monolith for 10,000 users.", source_id="smartlinx"),
                    TailoredHighlight(text="Led a team of 11 engineers.", source_id="smartlinx"),
                ],
            ),
            TailoredRole(
                id="pra-group", title="Systems Architect", company="PRA Group",
                date_range="Jan 2018 – Dec 2024", stack=["Oracle"],
                highlights=[TailoredHighlight(text="Cut annual cost from $2M to $17K.", source_id="pra-group")],
            ),
        ],
        skills=["C#", "Azure", "Terraform"],
        contact_email="shterzer@gmail.com",
        contact_linkedin="https://linkedin.com/in/talshterzer",
        profile_name="Tal Shterzer",
        profile_tagline="Engineering Manager",
    )


@pytest.fixture
def rendered_path():
    path = render_cv(_fixture_cv())
    yield path
    try:
        os.remove(path)
    except FileNotFoundError:
        pass


def _doc_text(path: str) -> str:
    return "\n".join(p.text for p in Document(path).paragraphs)


class TestTemplate:
    def test_template_is_bundled(self):
        assert TEMPLATE_PATH.exists()
        assert TEMPLATE_PATH.stat().st_size > 0


class TestRender:
    def test_writes_openable_nonempty_docx_in_tmp(self, rendered_path):
        p = Path(rendered_path)
        assert p.parent == Path("/tmp")
        assert p.suffix == ".docx"
        assert p.stat().st_size > 0
        Document(rendered_path)  # opens without raising => a real .docx

    def test_every_contract_field_reaches_the_document(self, rendered_path):
        text = _doc_text(rendered_path)
        assert "Tal Shterzer" in text                       # profile_name
        assert "Engineering Manager" in text                # profile_tagline
        assert "Staff Engineer" in text                     # target_role
        assert "shterzer@gmail.com" in text                 # contact_email
        assert "linkedin.com/in/talshterzer" in text        # contact_linkedin
        assert "Architect & manager with 17+ years." in text  # generated_summary
        assert "Terraform" in text                          # skills|join
        assert "SmartLinx" in text                          # role.company
        assert "Jan 2025 – Jun 2026" in text                # role.date_range
        assert "Re-platformed a monolith for 10,000 users." in text  # nested highlight
        assert "Oracle" in text                             # role.stack|join
        assert "PRA Group" in text                          # second role unrolled

    def test_special_chars_escaped_not_corrupting(self, rendered_path):
        text = _doc_text(rendered_path)
        assert "Hands-On Manager & Architect" in text  # '&' survived as valid XML
        assert "&amp;" not in text                     # and was not double-escaped

    def test_filenames_are_unique_per_render(self):
        a = render_cv(_fixture_cv())
        b = render_cv(_fixture_cv())
        try:
            assert a != b  # random token stem
        finally:
            for p in (a, b):
                try:
                    os.remove(p)
                except FileNotFoundError:
                    pass


class TestFullCvRender:
    def test_full_cv_renders_without_the_tailored_for_line(self):
        cv = build_full_cv(
            source_roles=[
                {"id": "a", "title": "Architect", "company": "SmartLinx",
                 "start": "2025-01", "end": "2026-06", "current": False,
                 "stack": ["C#"], "highlights": ["Did X for 10,000 users."]},
            ],
            profile={"name": "Tal Shterzer", "tagline": "Engineering Manager",
                     "summary": "17+ years."},
            contact={"email": "shterzer@gmail.com", "linkedin": "https://linkedin.com/in/talshterzer"},
            skills={"languages": ["C#"]},
        )
        path = render_cv(cv)
        try:
            text = _doc_text(path)
            assert "Tailored for" not in text          # conditional line omitted
            assert "Tal Shterzer" in text               # identity still renders
            assert "Did X for 10,000 users." in text    # highlight rendered
        finally:
            os.remove(path)


class TestPlaceholderDrift:
    def test_render_raises_when_a_template_field_is_missing_from_the_contract(self):
        """The regression guard itself: StrictUndefined makes a template placeholder
        with no matching model field fail loudly. Simulate drift by rendering a
        context missing a field the template references."""
        doc = DocxTemplate(str(TEMPLATE_PATH))
        env = jinja2.Environment(undefined=jinja2.StrictUndefined, autoescape=True)
        incomplete = _fixture_cv().model_dump()
        del incomplete["profile_name"]
        with pytest.raises(jinja2.UndefinedError):
            doc.render(incomplete, jinja_env=env)
